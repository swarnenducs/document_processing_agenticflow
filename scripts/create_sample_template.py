"""Create a sample Word invoice template with placeholders and styled runs."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def build_sample_template(output_path: Path) -> Path:
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("INVOICE")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    run.font.name = "Calibri"

    meta = doc.add_paragraph()
    r1 = meta.add_run("Invoice #: ")
    r1.bold = True
    r1.font.name = "Calibri"
    r1.font.size = Pt(12)
    r2 = meta.add_run("{{invoice_number}}")
    r2.font.name = "Calibri"
    r2.font.size = Pt(12)

    date_p = doc.add_paragraph()
    d1 = date_p.add_run("Date: ")
    d1.bold = True
    d1.font.name = "Calibri"
    d2 = date_p.add_run("{{invoice_date}}")
    d2.font.name = "Calibri"

    doc.add_paragraph()

    bill_to = doc.add_paragraph()
    bt = bill_to.add_run("Bill To")
    bt.bold = True
    bt.font.size = Pt(14)
    bt.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    bt.font.name = "Calibri"

    customer = doc.add_paragraph("{{customer.name}}")
    for run in customer.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(12)

    address = doc.add_paragraph("{{customer.address}}")
    for run in address.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.italic = True

    email = doc.add_paragraph("{{customer.email}}")
    for run in email.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)

    doc.add_paragraph()

    table = doc.add_table(rows=3, cols=4)
    table.style = "Table Grid"
    headers = ["Description", "Qty", "Unit Price", "Total"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Calibri"

    # Row 1 placeholders (first line item summary for the sample)
    table.rows[1].cells[0].text = "{{items[0].description}}"
    table.rows[1].cells[1].text = "{{items[0].quantity}}"
    table.rows[1].cells[2].text = "{{items[0].unit_price}}"
    table.rows[1].cells[3].text = "{{items[0].total}}"

    table.rows[2].cells[0].text = "{{items[1].description}}"
    table.rows[2].cells[1].text = "{{items[1].quantity}}"
    table.rows[2].cells[2].text = "{{items[1].unit_price}}"
    table.rows[2].cells[3].text = "{{items[1].total}}"

    doc.add_paragraph()

    totals = doc.add_paragraph()
    t1 = totals.add_run("Subtotal: ")
    t1.bold = True
    t1.font.name = "Calibri"
    t2 = totals.add_run("{{subtotal}}")
    t2.font.name = "Calibri"

    tax = doc.add_paragraph()
    x1 = tax.add_run("Tax: ")
    x1.bold = True
    x1.font.name = "Calibri"
    x2 = tax.add_run("{{tax}}")
    x2.font.name = "Calibri"

    total = doc.add_paragraph()
    g1 = total.add_run("Total Due: ")
    g1.bold = True
    g1.font.size = Pt(14)
    g1.font.name = "Calibri"
    g1.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    g2 = total.add_run("{{total_amount}}")
    g2.bold = True
    g2.font.size = Pt(14)
    g2.font.name = "Calibri"
    g2.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    notes = doc.add_paragraph()
    n1 = notes.add_run("Notes: ")
    n1.italic = True
    n1.font.name = "Calibri"
    n2 = notes.add_run("{{notes}}")
    n2.italic = True
    n2.font.name = "Calibri"
    n2.font.size = Pt(10)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    target = Path(__file__).resolve().parents[1] / "samples" / "templates" / "invoice_template.docx"
    path = build_sample_template(target)
    print(f"Wrote sample template → {path}")
