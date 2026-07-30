"""Voice → contract workflow with flexible parse + human-in-the-loop confirm."""

from __future__ import annotations

import re
from dataclasses import asdict, dataclass, field
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

_IRRELEVANT_MSG = "Please ask a relevant service."
_CONFIRM_HINTS = re.compile(
    r"^\s*(yes|y|ok|okay|confirm|proceed|create\s+it|go\s+ahead|select)\b",
    re.IGNORECASE,
)

_CREATE_CONTRACT_HINTS = re.compile(
    r"\b("
    r"create\s+(a\s+)?contract|"
    r"make\s+(a\s+)?contract|"
    r"new\s+contract|"
    r"generate\s+(a\s+)?contract|"
    r"draft\s+(a\s+)?contract|"
    r"prepare\s+(a\s+)?contract"
    r")\b",
    re.IGNORECASE,
)

# Flexible prompts, including missing "with" and spaced refs like "CR 1001"
_ENTITY_REF_PATTERNS: list[re.Pattern[str]] = [
    re.compile(
        r"(?:please\s+)?create\s+(?:a\s+)?contract\s+with\s+legal\s+entity\s+"
        r"(?P<entity>.+?)\s+(?:with\s+)?contract\s+reference\s+number\s+"
        r"(?P<ref>[A-Za-z0-9][\w\-\s]*[A-Za-z0-9]|[A-Za-z0-9]+)",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:please\s+)?create\s+(?:a\s+)?contract\s+for\s+legal\s+entity\s+"
        r"(?P<entity>.+?)\s+(?:with\s+)?(?:contract\s+)?(?:reference\s+)?number\s+"
        r"(?P<ref>[A-Za-z0-9][\w\-\s]*[A-Za-z0-9]|[A-Za-z0-9]+)",
        re.IGNORECASE,
    ),
    re.compile(
        r"legal\s+entity\s+(?P<entity>.+?)\s+"
        r"(?:with\s+)?contract\s+reference\s+number\s+"
        r"(?P<ref>[A-Za-z0-9][\w\-\s]*[A-Za-z0-9]|[A-Za-z0-9]+)",
        re.IGNORECASE,
    ),
]


@dataclass
class VoiceContractResult:
    ok: bool
    message: str
    intent: str | None = None
    status: str = "rejected"  # rejected | needs_confirmation | completed
    legal_entity_name: str | None = None
    contract_reference_number: str | None = None
    legal_entity: dict[str, Any] | None = None
    pricelist: dict[str, Any] | None = None
    candidates: list[dict[str, Any]] = field(default_factory=list)
    contract_payload: dict[str, Any] | None = None
    contract_file: str | None = None
    contract_text_file: str | None = None
    contract_text: str | None = None
    transcript: str | None = None
    contract_id: str | None = None
    thread_id: str | None = None  # LangGraph HITL thread
    spoken_name: str | None = None
    spoken_number: str | None = None
    contact: dict[str, Any] | None = None

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


def _normalize(value: str) -> str:
    cleaned = re.sub(r"[^\w\s\-]", " ", (value or "").lower())
    return re.sub(r"\s+", " ", cleaned).strip()


def normalize_contract_ref(value: str) -> str:
    """Normalize spoken/typed refs: 'CR 1001' / 'cr-1001' → 'CR1001' for matching."""
    return re.sub(r"[\s\-_]+", "", (value or "").upper())


def format_contract_ref(value: str) -> str:
    """Pretty-print refs like CR1001 → CR-1001 when pattern matches."""
    raw = (value or "").strip().upper()
    compact = normalize_contract_ref(raw)
    match = re.fullmatch(r"([A-Z]+)(\d+)", compact)
    if match:
        return f"{match.group(1)}-{match.group(2)}"
    return raw or compact


def extract_legal_entity_and_reference(text: str) -> tuple[str | None, str | None]:
    cleaned = " ".join((text or "").strip().split())
    if not cleaned:
        return None, None
    for pattern in _ENTITY_REF_PATTERNS:
        match = pattern.search(cleaned)
        if not match:
            continue
        entity = " ".join(match.group("entity").split()).strip(" .,;:")
        entity = re.sub(
            r"\b(with|and|the|a|an|please|for|contract|reference|number)\b$",
            "",
            entity,
            flags=re.IGNORECASE,
        ).strip(" .,;:")
        ref = " ".join(match.group("ref").split()).strip(" .,;:")
        if entity and ref:
            return entity, ref
    return None, None


def _is_create_contract_intent(text: str) -> bool:
    return bool(_CREATE_CONTRACT_HINTS.search(text or ""))


def is_confirmation(text: str) -> bool:
    return bool(_CONFIRM_HINTS.search(text or ""))


def build_contract_payload(
    entity: dict[str, Any],
    pricelist: dict[str, Any],
    *,
    contract_reference_number: str,
) -> dict[str, Any]:
    ref = format_contract_ref(
        str(pricelist.get("contractReferenceNumber") or contract_reference_number)
    )
    return {
        "DATE": date.today().isoformat(),
        "accountName": entity.get("legalName") or entity.get("code"),
        "legalEntityCode": entity.get("code"),
        "legalName": entity.get("legalName"),
        "address": entity.get("address"),
        "city": entity.get("city"),
        "country": entity.get("country"),
        "postalCode": entity.get("postalCode"),
        "email": entity.get("email"),
        "phone": entity.get("phone"),
        "registrationNumber": entity.get("registrationNumber"),
        "agreementNumber": ref,
        "contractReferenceNumber": ref,
        "currency": pricelist.get("currency") or "USD",
        "effectiveDate": pricelist.get("effectiveDate"),
        "products": list(pricelist.get("products") or []),
    }


def render_contract_text(payload: dict[str, Any]) -> str:
    products = payload.get("products") or []
    lines = [
        "SUPPLY CONTRACT (DUMMY)",
        "=" * 48,
        f"Contract Reference Number : {payload.get('contractReferenceNumber')}",
        f"Date                      : {payload.get('DATE')}",
        "",
        "LEGAL ENTITY",
        "-" * 48,
        f"Legal Name   : {payload.get('legalName')}",
        f"Code         : {payload.get('legalEntityCode')}",
        f"Address      : {payload.get('address')}",
        f"City         : {payload.get('city')}",
        f"Country      : {payload.get('country')}",
        f"Postal Code  : {payload.get('postalCode')}",
        f"Email        : {payload.get('email')}",
        f"Phone        : {payload.get('phone')}",
        f"Registration : {payload.get('registrationNumber')}",
        "",
        "PRICELIST / PRODUCTS",
        "-" * 48,
        f"Currency       : {payload.get('currency')}",
        f"Effective Date : {payload.get('effectiveDate')}",
        "",
        f"{'Code':<12} {'Description':<40} {'UOM':<6} {'Price':>8} {'Market':>8}",
    ]
    for product in products:
        lines.append(
            f"{str(product.get('productCode') or ''):<12} "
            f"{str(product.get('productDescription') or '')[:40]:<40} "
            f"{str(product.get('uom') or ''):<6} "
            f"{str(product.get('unitPrice') or ''):>8} "
            f"{str(product.get('marketPrice') or ''):>8}"
        )
    lines.extend(
        [
            "",
            "This is a dummy contract generated after human confirmation "
            "from SQLite legal-entity and pricelist data.",
        ]
    )
    return "\n".join(lines)


def create_dummy_contract_docx(payload: dict[str, Any], output_path: Path) -> Path:
    """Create a simple Word contract from legal entity + pricelist (no LLM)."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SUPPLY CONTRACT (DUMMY)")
    run.bold = True
    run.font.size = Pt(18)

    meta = doc.add_paragraph()
    meta.add_run("Contract Reference Number: ").bold = True
    meta.add_run(str(payload.get("contractReferenceNumber") or ""))

    date_p = doc.add_paragraph()
    date_p.add_run("Date: ").bold = True
    date_p.add_run(str(payload.get("DATE") or ""))

    doc.add_heading("Legal Entity", level=2)
    for label, key in (
        ("Legal Name", "legalName"),
        ("Code", "legalEntityCode"),
        ("Address", "address"),
        ("City", "city"),
        ("Country", "country"),
        ("Postal Code", "postalCode"),
        ("Email", "email"),
        ("Phone", "phone"),
        ("Registration No.", "registrationNumber"),
    ):
        line = doc.add_paragraph()
        line.add_run(f"{label}: ").bold = True
        line.add_run(str(payload.get(key) or "—"))

    products = payload.get("products") or []
    doc.add_heading("Pricelist / Products", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Product Code", "Description", "UOM", "Unit Price", "Market Price"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header

    for product in products:
        row = table.add_row().cells
        row[0].text = str(product.get("productCode") or "")
        row[1].text = str(product.get("productDescription") or "")
        row[2].text = str(product.get("uom") or "")
        row[3].text = str(product.get("unitPrice") or product.get("eaPrice") or "")
        row[4].text = str(product.get("marketPrice") or "")

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run(
        "This is a dummy contract generated after human confirmation "
        "from SQLite legal-entity and pricelist data."
    ).italic = True

    doc.save(str(output_path))
    return output_path


def create_dummy_contract_files(
    payload: dict[str, Any],
    output_dir: Path,
    *,
    stem: str,
) -> tuple[Path, Path, str]:
    output_dir.mkdir(parents=True, exist_ok=True)
    text = render_contract_text(payload)
    text_path = output_dir / f"{stem}.txt"
    docx_path = output_dir / f"{stem}.docx"
    text_path.write_text(text, encoding="utf-8")
    create_dummy_contract_docx(payload, docx_path)
    return docx_path, text_path, text


def finalize_contract(
    *,
    entity: dict[str, Any],
    pricelist: dict[str, Any],
    store: Any,
    transcript: str | None = None,
    output_dir: Path | None = None,
) -> VoiceContractResult:
    ref = format_contract_ref(str(pricelist.get("contractReferenceNumber") or ""))
    payload = build_contract_payload(entity, pricelist, contract_reference_number=ref)
    if output_dir is None:
        output_dir = Path(store.cfg.storage_base_path) / "voice_contracts"
    stem = f"draft_{_normalize(str(entity.get('code') or 'entity'))}_{normalize_contract_ref(ref)}"
    docx_path, text_path, text = create_dummy_contract_files(
        payload, Path(output_dir), stem=stem
    )
    return VoiceContractResult(
        ok=True,
        status="completed",
        message=(
            f"Contract confirmed and created for legal entity "
            f"'{entity.get('legalName')}' using reference '{ref}'."
        ),
        intent="create_contract",
        legal_entity_name=str(entity.get("code") or entity.get("legalName")),
        contract_reference_number=ref,
        legal_entity=entity,
        pricelist=pricelist,
        candidates=[pricelist],
        contract_payload=payload,
        contract_file=str(docx_path),
        contract_text_file=str(text_path),
        contract_text=text,
        spoken_name=str(entity.get("code") or entity.get("legalName")),
        spoken_number=ref,
        contact=entity,
        transcript=transcript,
    )


def run_voice_contract_workflow(
    transcript: str,
    *,
    store: Any | None = None,
    output_dir: Path | None = None,
    auto_create: bool = False,
    thread_id: str | None = None,
) -> VoiceContractResult:
    """
    Run the LangGraph voice-contract agent.

    Data lookups (legal entity / pricelist) happen in graph nodes via SQLite
    (swappable later for HTTP APIs). HITL confirmation uses LangGraph interrupt().
    """
    del store  # lookups use JobStore inside graph nodes
    from document_processing_agenticflow.voice_graph import start_voice_contract_agent

    result, tid = start_voice_contract_agent(
        transcript,
        auto_create=auto_create,
        thread_id=thread_id,
        output_dir=str(output_dir) if output_dir else None,
    )
    result.thread_id = tid
    return result


def confirm_voice_contract(
    *,
    entity_code_or_name: str,
    contract_reference_number: str,
    store: Any | None = None,
    transcript: str | None = None,
    output_dir: Path | None = None,
    thread_id: str | None = None,
    user_text: str | None = None,
) -> VoiceContractResult:
    """
    Human-in-the-loop confirmation.

    Prefer ``thread_id`` to resume the LangGraph interrupt. If no thread is
    available, fall back to a direct finalize path (compatibility).
    """
    del store
    if thread_id:
        from document_processing_agenticflow.voice_graph import resume_voice_contract_agent

        result = resume_voice_contract_agent(
            thread_id,
            user_text=user_text or "yes",
            action=user_text or "yes",
            contract_reference_number=contract_reference_number,
        )
        result.thread_id = thread_id
        return result

    # Compatibility fallback (no active LangGraph thread)
    from document_processing_agenticflow.storage.job_store import JobStore

    job_store = JobStore()
    job_store.ensure_contract_catalog_seeded()
    entity = job_store.find_legal_entity(entity_code_or_name)
    if entity is None:
        return VoiceContractResult(
            ok=False,
            status="rejected",
            message=f"No legal entity found for '{entity_code_or_name}'.",
            intent="create_contract",
        )
    matches = job_store.search_pricelists(
        contract_reference_number,
        legal_entity_code=str(entity.get("code") or "") or None,
    )
    if not matches:
        matches = job_store.search_pricelists(contract_reference_number)
    if not matches:
        return VoiceContractResult(
            ok=False,
            status="rejected",
            message=(
                f"No pricelist found for reference '{contract_reference_number}'. "
                f"{_IRRELEVANT_MSG}"
            ),
            intent="create_contract",
            legal_entity=entity,
            contact=entity,
        )
    result = finalize_contract(
        entity=entity,
        pricelist=matches[0],
        store=job_store,
        transcript=transcript,
        output_dir=output_dir,
    )
    return result
